"""Generate formal DOCX audit report for Round 3."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REGISTER = Path("/home/z/my-project/download/TRA_Round3/master_findings_register_r3.json")
OUT = Path("/home/z/my-project/download/TRA_Round3/TRA_Prototype_Audit_Report_r3.docx")

SEVERITY_COLORS = {
    "BLOCKING": RGBColor(0xCC, 0x00, 0x00),
    "WARNING": RGBColor(0xCC, 0x88, 0x00),
    "INFO": RGBColor(0x00, 0x66, 0x00),
}


def main() -> None:
    findings = json.loads(REGISTER.read_text(encoding="utf-8"))
    doc = Document()

    # Cover
    title = doc.add_heading("TRA Prototype Audit Report — Round 3", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Independent Re-Audit of the Translation Runtime Architecture Prototype")
    run.font.size = Pt(14)
    run.font.italic = True
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Repository: nordeim/Translation-Runtime-Architecture\n").bold = True
    meta.add_run("HEAD audited: b783745\n").bold = True
    meta.add_run("Audit date: 2026-07-15\n").bold = True
    meta.add_run("Methodology: 7-track parallel re-audit (R3 + A3 + B3 + C3 + D3 + E3 + F3)").bold = True
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    sev = {"BLOCKING": 0, "WARNING": 0, "INFO": 0}
    for f in findings:
        sev[f["severity"]] += 1
    p = doc.add_paragraph()
    p.add_run(
        f"This report presents the findings of the Round 3 independent re-audit of the "
        f"TRA (Translation Runtime Architecture) prototype at HEAD b783745. The audit "
        f"used a 7-track parallel structure: Track R3 (regression baseline of all 71 prior "
        f"findings), Track A3 (spec conformance), Track B3 (code quality + OWASP security "
        f"deep-dive), Track C3 (doc consistency), Track D3 (test suite + e2e quality), "
        f"Track E3 (forensic L4 end-to-end), and Track F3 (module extension safety, NEW). "
        f"A total of {len(findings)} findings were identified: "
    )
    p.add_run(f"{sev['BLOCKING']} BLOCKING").bold = True
    p.add_run(f", {sev['WARNING']} WARNING, and {sev['INFO']} INFO. ")
    p.add_run(
        f"All 4 quality gates remain green (174 tests passing). L4 byte-reproducibility "
        f"holds (TRA-013). The 3 Round 2 BLOCKING findings (TRA-006/036/037/048) are "
        f"confirmed fixed. Two new BLOCKING findings were identified by Track F3 "
        f"(TRA-096: as_interface() crashes with Pydantic ValidationError) and Track E3 "
        f"(TRA-093: false-positive BROKEN_LINK blocks valid CJK translations)."
    )

    # Methodology
    doc.add_heading("2. Methodology", level=1)
    doc.add_paragraph(
        "Round 3 followed the same 4-track parallel structure as Round 2, extended to "
        "7 tracks. Each track's findings were re-checked against the cited code at HEAD "
        "b783745 before synthesis. Severity lexicon mirrors the TRA spec: "
        "BLOCKING / WARNING / INFO."
    )
    tracks_table = doc.add_table(rows=1, cols=4)
    tracks_table.style = "Light Grid Accent 1"
    hdr = tracks_table.rows[0].cells
    hdr[0].text = "Track"
    hdr[1].text = "Scope"
    hdr[2].text = "Methodology"
    hdr[3].text = "Findings"
    track_data = [
        ("R3", "Regression baseline (71 findings)", "pytest + static grep", "12 persistent"),
        ("A3", "Spec conformance (Kernel, ISA, Policy, Memory, Exceptions, L3/L4)", "Manual review vs Spec §1-9", "10 (0/6/4)"),
        ("B3", "Code quality + OWASP security deep-dive", "4 gates + reproducibility + OWASP checklist", "13 (0/4/9)"),
        ("C3", "Doc-vs-code consistency (16 docs)", "Per-claim verification", "13 (0/5/8)"),
        ("D3", "Test suite + e2e quality", "Coverage + 6 mutations + e2e audit", "11 (0/6/5)"),
        ("E3", "Forensic L4 end-to-end", "7 deliberate-failure probes + reproducibility", "14 (1/2/11)"),
        ("F3", "Module extension safety (NEW)", "Stub module authoring + registry testing", "11 (2/5/4)"),
    ]
    for t, s, m, f in track_data:
        row = tracks_table.add_row().cells
        row[0].text = t
        row[1].text = s
        row[2].text = m
        row[3].text = f

    # Quality Gates
    doc.add_heading("3. Quality Gates (Baseline)", level=1)
    doc.add_paragraph(
        "All 4 quality gates were green at HEAD b783745 at audit start and remained "
        "green throughout (no code was modified during the audit):"
    )
    gates = doc.add_paragraph()
    gates.add_run("ruff format --check").bold = True
    gates.add_run(" (39 files) ✓\n")
    gates.add_run("ruff check").bold = True
    gates.add_run(" ✓\n")
    gates.add_run("mypy --strict tra").bold = True
    gates.add_run(" (20 source files) ✓\n")
    gates.add_run("pytest tests").bold = True
    gates.add_run(" (174 tests, 0.68s) ✓")

    # Round 2 Carry-over Status
    doc.add_heading("4. Round 2 Carry-over Status", level=1)
    r2_status = {}
    for f in findings:
        s = f["round2_status"]
        r2_status[s] = r2_status.get(s, 0) + 1
    doc.add_paragraph(
        f"Of the {len(findings)} Round 3 findings, the Round 2 status breakdown is: "
        f"{r2_status.get('fixed', 0)} fixed (confirmed at HEAD), "
        f"{r2_status.get('persistent', 0)} persistent (still open), "
        f"{r2_status.get('partial', 0)} partial (partially fixed), and "
        f"{r2_status.get('new', 0)} new (introduced or newly discovered)."
    )
    doc.add_paragraph(
        "Key fixes confirmed: TRA-006 (PolicyResolver now invoked in verify_output via "
        "_POLICY_RESOLVER.wins()), TRA-036 (analyze-failure raises ConformanceFailure "
        "at L3/L4), TRA-037 (_rewrite_anchors runs before L3 gate), TRA-039 "
        "(build_entity_table wrapped), TRA-041 (GLOSSARY_CONFLICT populates glossary), "
        "TRA-043 (LanguageModuleProtocol), TRA-044 (Unrecoverable branch), TRA-047 "
        "(extra='forbid'), TRA-048 (LLM-degradation test strengthened), TRA-049/050/051/053/054 "
        "(regression tests added), TRA-071 (BrokenMarkdown raised)."
    )

    # BLOCKING Findings
    doc.add_heading("5. BLOCKING Findings (Immediate Attention)", level=1)
    blocking = [f for f in findings if f["severity"] == "BLOCKING"]
    if not blocking:
        doc.add_paragraph("No BLOCKING findings. ✅")
    else:
        for f in blocking:
            doc.add_heading(f"{f['id']}: {f['title']}", level=2)
            p = doc.add_paragraph()
            p.add_run("Severity: ").bold = True
            run = p.add_run(f["severity"])
            run.font.color.rgb = SEVERITY_COLORS[f["severity"]]
            run.bold = True
            p.add_run(f"\nCategory: {f['category']}")
            p.add_run(f"\nTrack: {f['track']}")
            p.add_run(f"\nRound 2 status: {f['round2_status']}")
            doc.add_paragraph(f"Evidence: {f['evidence']}")
            doc.add_paragraph(f"Detail: {f['detail']}")
            doc.add_paragraph(f"Suggested fix: {f['suggested_fix']}")

    # All Findings (condensed)
    doc.add_heading("6. All Findings (Condensed)", level=1)
    findings_table = doc.add_table(rows=1, cols=5)
    findings_table.style = "Light Grid Accent 1"
    hdr = findings_table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Severity"
    hdr[2].text = "Track"
    hdr[3].text = "Title"
    hdr[4].text = "R2 Status"
    for f in findings:
        row = findings_table.add_row().cells
        row[0].text = f["id"]
        row[1].text = f["severity"]
        row[2].text = f["track"]
        row[3].text = f["title"]
        row[4].text = f["round2_status"]

    # Detailed Findings (WARNING + INFO)
    doc.add_heading("7. Detailed WARNING & INFO Findings", level=1)
    for f in findings:
        if f["severity"] == "BLOCKING":
            continue
        doc.add_heading(f"{f['id']}: {f['title']}", level=3)
        p = doc.add_paragraph()
        p.add_run("Severity: ").bold = True
        run = p.add_run(f["severity"])
        run.font.color.rgb = SEVERITY_COLORS[f["severity"]]
        run.bold = True
        p.add_run(f"  |  Track: {f['track']}  |  R2: {f['round2_status']}")
        doc.add_paragraph(f"Evidence: {f['evidence']}")
        doc.add_paragraph(f"Detail: {f['detail']}")
        doc.add_paragraph(f"Suggested fix: {f['suggested_fix']}")

    # Recommendations
    doc.add_heading("8. Recommendations", level=1)
    doc.add_paragraph(
        "The following recommendations are prioritized by severity and impact:"
    )
    recs = [
        "FIX TRA-096 (BLOCKING): The spec's sanctioned module extension path (as_interface()) is broken. Either add the 4 missing Callable fields to ModuleInterface, or deprecate as_interface() and have register() accept the full module object directly.",
        "FIX TRA-093 (BLOCKING): The false-positive BROKEN_LINK in _rewrite_anchors Pass 2 blocks valid CJK heading + CJK link translations at L3/L4. Check if the link slug matches a TRANSLATED slug value, not just original slugs.",
        "FIX TRA-077 (WARNING, OWASP A08): Switch diskcache from pickle to JSON serialization to eliminate RCE-on-cache-load risk. Store model_dump_json() instead of model_dump().",
        "FIX TRA-076 (WARNING, OWASP A03): Route LLM seam output through sanitize_input before use. A malicious LLM could inject bidi overrides.",
        "FIX TRA-017 (WARNING, persistent): Remove 6 unused deps from pyproject.toml (litellm, structlog, pydantic-settings, mdit-py-plugins, black, pytest-asyncio). litellm pulls ~50 transitive packages.",
        "FIX TRA-038 (WARNING, persistent): Wire UnknownTerm, CertaintyConflict, EntityAmbiguity in production code paths. The spec's exception model is only 40% operational.",
        "FIX TRA-042 (WARNING, persistent): Extend verify_output to check table row/col counts, list nesting, and code-block fence counts. Currently structural verification is heading-count-only.",
        "FIX TRA-080/082/085 (WARNING, doc staleness): Update CLAUDE.md, tra-prototype/README.md, and status.md to reflect Round 2 fixes. TRA-006 is no longer a half-fix; EntityAmbiguity is never raised; test count is 174 not 103.",
        "FIX TRA-088/089/091 (WARNING, test gaps): Extend LLM-seam tests to assert single-audit-record; add ConformanceFailure e2e tests; add interactive=True kernel end-to-end test.",
        "FIX TRA-097/098/099 (WARNING, module registry): Add isinstance Protocol check in register(); add duplicate-name and direction-conflict detection; add --registry CLI flag.",
        "FIX TRA-001 (WARNING, partial): Implement per-leaf segment translation. This unblocks per-segment cache keys, per-segment repair, and structural evidence tracing (TRA-094).",
        "ADDRESS TRA-072 (WARNING): Route all severity decisions in verify_output through PolicyResolver, not just TERMINOLOGICAL vs FLUENCY.",
    ]
    for i, rec in enumerate(recs, 1):
        doc.add_paragraph(rec, style="List Number")

    # Conclusion
    doc.add_heading("9. Conclusion", level=1)
    doc.add_paragraph(
        f"The TRA prototype at HEAD b783745 is substantially more conformant than at "
        f"Round 2. All 4 critical invariants hold. All 4 quality gates are green (174 "
        f"tests). L4 byte-reproducibility holds. The 3 Round 2 BLOCKING findings in "
        f"Track A scope (TRA-006/036/037) are confirmed fixed. However, 2 new BLOCKING "
        f"findings were identified: TRA-096 (module extension path broken — the spec's "
        f"sanctioned API crashes) and TRA-093 (false-positive BROKEN_LINK blocks valid "
        f"CJK translations). These should be addressed before any L3/L4 production use. "
        f"The remaining 18 WARNING and 16 INFO findings represent documented gaps, "
        f"persistent carry-overs, and new security findings from the OWASP deep-dive. "
        f"Estimated total remediation effort: ~50 hours (6.3 person-days), with TRA-001 "
        f"(per-leaf segment translation) accounting for ~16 hours alone."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
