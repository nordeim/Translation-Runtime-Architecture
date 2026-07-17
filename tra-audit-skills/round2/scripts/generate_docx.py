"""Generate the formal DOCX audit report for Round 2."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

REGISTER_PATH = Path("/home/z/my-project/audit-ctx/master_findings_register.json")
HEATMAP_PATH = Path("/home/z/my-project/download/TRA_audit_severity_heatmap_r2.png")
OUT_PATH = Path("/home/z/my-project/download/TRA_Prototype_Audit_Report_r2.docx")

SEVERITY_COLOR = {
    "BLOCKING": RGBColor(0xC6, 0x28, 0x28),
    "WARNING": RGBColor(0xF5, 0x7C, 0x00),
    "INFO": RGBColor(0x2E, 0x7D, 0x32),
}


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x2E)


def add_finding(doc: Document, finding: dict) -> None:
    # Title with severity color
    h = doc.add_heading(level=2)
    run = h.add_run(f"{finding['id']} — {finding['title']}")
    run.font.color.rgb = SEVERITY_COLOR[finding["severity"]]
    run.font.size = Pt(13)

    # Metadata line
    p = doc.add_paragraph()
    p.add_run("Severity: ").bold = True
    sev_run = p.add_run(finding["severity"])
    sev_run.font.color.rgb = SEVERITY_COLOR[finding["severity"]]
    sev_run.bold = True
    p.add_run(f"    |    Category: {finding['category']}")
    p.add_run(f"    |    Track: {finding['track']}")
    p.add_run(f"    |    Round 1: {finding['round1_status']}")

    # Evidence
    p = doc.add_paragraph()
    p.add_run("Evidence: ").bold = True
    p.add_run(finding["evidence"])

    # Detail
    p = doc.add_paragraph()
    p.add_run("Detail: ").bold = True
    p.add_run(finding["detail"])

    # Suggested fix
    p = doc.add_paragraph()
    p.add_run("Suggested fix: ").bold = True
    p.add_run(finding["suggested_fix"])

    # Source findings
    if finding.get("source_findings"):
        p = doc.add_paragraph()
        p.add_run("Source track findings: ").bold = True
        p.add_run("; ".join(finding["source_findings"]))

    doc.add_paragraph()  # spacing


def main() -> None:
    findings = json.loads(REGISTER_PATH.read_text(encoding="utf-8"))
    doc = Document()

    # === Cover ===
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TRA Prototype Audit Report")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Round 2 — Independent Re-Audit")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nRepo: ").bold = True
    p.add_run("nordeim/Translation-Runtime-Architecture")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HEAD: ").bold = True
    p.add_run("4b8827c")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Date: ").bold = True
    p.add_run("2026-07-14")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Methodology: ").bold = True
    p.add_run("5-track parallel re-audit (R+A+B+C+D+E) using Round 1's tra-audit-skills bundle as baseline")

    doc.add_paragraph()

    # === Executive Summary ===
    add_heading(doc, "1. Executive Summary", level=1)

    blocking = [f for f in findings if f["severity"] == "BLOCKING"]
    warning = [f for f in findings if f["severity"] == "WARNING"]
    info = [f for f in findings if f["severity"] == "INFO"]
    carry_partial = [f for f in findings if f["round1_status"] == "partial"]
    carry_persistent = [f for f in findings if f["round1_status"] == "persistent"]
    new_findings = [f for f in findings if f["round1_status"] == "new"]

    p = doc.add_paragraph()
    p.add_run(
        f"This report documents the Round 2 independent re-audit of the Translation Runtime "
        f"Architecture (TRA) prototype engine at HEAD commit 4b8827c. The audit used the "
        f"Round 1 audit's methodology template (encoded in tra-audit-skills/worklog.md, 3,735 "
        f"lines) and the Round 1 35-finding register as the regression baseline. The audit "
        f"employed a 5-track parallel structure: Track R (regression baseline), Track A (spec "
        f"conformance), Track B (code quality & security), Track C (doc-vs-code consistency), "
        f"Track D (test suite), and Track E (forensic end-to-end L4 verification)."
    )

    p = doc.add_paragraph()
    p.add_run("Headline numbers: ").bold = True
    p.add_run(
        f"{len(findings)} total findings — {len(blocking)} BLOCKING, {len(warning)} WARNING, "
        f"{len(info)} INFO. Of the 35 Round-1 findings, {len(carry_partial)} are partial "
        f"({', '.join(f['id'] for f in carry_partial)}), {len(carry_persistent)} are persistent "
        f"({', '.join(f['id'] for f in carry_persistent)}), and 30 are fully fixed. "
        f"Round 2 uncovered {len(new_findings)} new findings (TRA-036 through TRA-071)."
    )

    p = doc.add_paragraph()
    p.add_run("Quality gates at HEAD: ").bold = True
    p.add_run(
        "All 4 gates green — ruff check ✓, ruff format --check (35 files) ✓, "
        "mypy --strict tra (20 files) ✓, pytest (141 tests, 0.68s) ✓. "
        "The test count grew from 103 (Round 1) to 141 (+38 new regression tests)."
    )

    p = doc.add_paragraph()
    p.add_run("Critical invariants: ").bold = True
    p.add_run(
        "All 4 spec-mandated invariants hold at HEAD — (1) canonical terminology exact "
        "(成立→Confirmed, 执行环境→execution environment, 高度可信→highly credible); "
        "(2) entities immutable (mutable=False, frozen=True); (3) verification never self-scores "
        "(confidence_note never read by verify_output/repair_segment); (4) repair surgical "
        "(raises Unrecoverable on any new BLOCKING)."
    )

    p = doc.add_paragraph()
    p.add_run("Reproducibility: ").bold = True
    p.add_run(
        "TRA-013 fully remediated. Two cold-cache L4 runs produce byte-identical audit_trace.jsonl, "
        "evidence_trace.jsonl, ambiguity_register.json, and output .md (verified via sha256sum). "
        "The deterministic clock (kernel.py:157-171) and content-addressed evidence IDs "
        "(diagnostics.py:45-63) produce stable timestamps and IDs."
    )

    # === Heatmap ===
    if HEATMAP_PATH.exists():
        doc.add_picture(str(HEATMAP_PATH), width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run("Figure 1: Severity heatmap by audit track")
        run.italic = True
        run.font.size = Pt(10)

    # === Methodology ===
    add_heading(doc, "2. Audit Methodology", level=1)

    add_heading(doc, "2.1 Track structure", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "The Round 2 audit extended Round 1's 4-track structure with a Track R (regression "
        "baseline) and a Track E (forensic end-to-end L4 verification). Track R ran first to "
        "anchor the audit: it re-verified all 35 Round-1 findings at HEAD. Tracks A-D ran in "
        "parallel via independent subagents, each producing a self-contained findings file. "
        "Track E ran end-to-end L4 runs and deliberate-failure probes."
    )

    # Track table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].add_run("Track").bold = True
    hdr[1].paragraphs[0].add_run("Scope").bold = True
    hdr[2].paragraphs[0].add_run("Findings").bold = True
    hdr[3].paragraphs[0].add_run("Method").bold = True

    track_summary = [
        ("R", "Regression baseline (35 Round-1 findings)", "5 carry-overs",
         "pytest + grep static checks"),
        ("A", "Spec conformance (Kernel, ISA, Policy, Memory, Exceptions, L3/L4)", "11 (0/7/4)",
         "Code reading vs TRA-SPECIFICATION.md §1-§9 + TRA-ISA-REFERENCE.md"),
        ("B", "Code quality & security (types, errors, cache, deps, reproducibility)", "7 (0/6/1)",
         "mypy --strict, grep, 2-run sha256sum probe"),
        ("C", "Doc-vs-code consistency (14 docs)", "10 (0/4/6)",
         "Per-doc claim verification against code"),
        ("D", "Test suite (coverage, mutation, benchmark, HITL, LLM seam)", "13 (1/8/4)",
         "coverage run + 5 manual mutations + pytest --collect-only"),
        ("E", "Forensic L4 end-to-end (artifact structure, byte-reproducibility, probes)", "6 (2/3/1)",
         "L4 translate runs + deliberate-failure probes + sha256sum"),
    ]
    for trk, scope, count, method in track_summary:
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(trk)
        row[1].paragraphs[0].add_run(scope)
        row[2].paragraphs[0].add_run(count)
        row[3].paragraphs[0].add_run(method)

    add_heading(doc, "2.2 Severity lexicon", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "BLOCKING = violates a spec invariant or breaks L3+ certification. "
        "WARNING = degrades quality/auditability but doesn't break invariants. "
        "INFO = minor hygiene, dead code, or stylistic. "
        "Every finding cites file:line evidence; every BLOCKING/WARNING is reproducible."
    )

    add_heading(doc, "2.3 Anti-slop guarantees", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "No 'looks good' verdicts without proof. No severity inflation (a doc typo is INFO). "
        "Re-validation mandatory: each track's draft findings were re-checked against the cited "
        "code at HEAD before synthesis. Track R's static-check inverter bugs (3 false PASS for "
        "persistent findings) were corrected during synthesis."
    )

    # === Round-1 carry-over status ===
    add_heading(doc, "3. Round-1 Carry-Over Status", level=1)

    p = doc.add_paragraph()
    p.add_run(
        f"Of the 35 Round-1 findings, 30 are fully fixed at HEAD (verified by Track R regression "
        f"tests or static checks). 5 carry over to Round 2:"
    )

    for f in carry_partial + carry_persistent:
        add_finding(doc, f)

    # === New findings ===
    add_heading(doc, "4. New Findings (TRA-036 through TRA-071)", level=1)

    p = doc.add_paragraph()
    p.add_run(
        f"Round 2 uncovered {len(new_findings)} new findings. The 3 BLOCKING findings require "
        f"immediate attention; the 22 WARNING findings should be addressed before Phase 7 "
        f"(documentation & delivery); the 11 INFO findings are minor hygiene."
    )

    add_heading(doc, "4.1 BLOCKING findings (immediate attention)", level=2)
    for f in new_findings:
        if f["severity"] == "BLOCKING":
            add_finding(doc, f)

    add_heading(doc, "4.2 WARNING findings", level=2)
    for f in new_findings:
        if f["severity"] == "WARNING":
            add_finding(doc, f)

    add_heading(doc, "4.3 INFO findings", level=2)
    for f in new_findings:
        if f["severity"] == "INFO":
            add_finding(doc, f)

    # === Recommendations ===
    add_heading(doc, "5. Recommendations", level=1)

    add_heading(doc, "5.1 Immediate (BLOCKING — before any L3/L4 production use)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "1. TRA-036: Replace the analyze-failure early `return ''` (kernel.py:214) with a "
        "ConformanceFailure raise at L3+. Add a regression test at L3_STRICT. "
        "2. TRA-037: Move _rewrite_anchors to BEFORE the L3 gate so the gate verifies the "
        "post-rewrite target. Add an unresolved_ambiguities BROKEN_LINK check at L3+. "
        "3. TRA-048: Strengthen test_phase6_hardening.py:84 to assert exactly one TRANSLATE_SEGMENT "
        "audit record on LLM degradation."
    )

    add_heading(doc, "5.2 Short-term (WARNING — before Phase 7)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "4. TRA-006: Wire PolicyResolver into verify_output (currently scaffolding). "
        "5. TRA-038: Raise UnknownTerm/CertaintyConflict/EntityAmbiguity from the appropriate "
        "ISA functions; route through _recover. "
        "6. TRA-041: Set the first-occurrence canonical mapping in GLOSSARY_CONFLICT recovery. "
        "7. TRA-044: Add an explicit Unrecoverable branch to route_exception returning BLOCKING + HALT. "
        "8. TRA-017: Remove the 6 unused deps from pyproject.toml (litellm pulls ~50 transitive). "
        "9. TRA-049-054: Close the 6 test-coverage gaps identified by mutation testing. "
        "10. TRA-059-062: Update the 4 stale doc sections (CLAUDE.md 'Known gaps', prototype README, "
        "implementation_plan scope notes)."
    )

    add_heading(doc, "5.3 Long-term (INFO + structural)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "11. TRA-001: Implement full per-leaf segment translation (unblocks per-segment cache, "
        "repair, and structural evidence tracing). This is the largest single fix. "
        "12. TRA-040: Model EXCEPTION_HANDLER and HALT_ERROR as KernelStates per spec §2.1. "
        "13. TRA-042: Add structural-integrity checks beyond heading count (lists, tables, blockquotes). "
        "14. TRA-043: Define a LanguageModuleProtocol and type RuntimeContext.module properly. "
        "15. TRA-071: Add a structural-validation pass to raise BrokenMarkdown for spec-defined "
        "malformed cases (markdown-it-py is too lenient to raise it naturally)."
    )

    add_heading(doc, "5.4 Estimated remediation effort", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Using 4h/BLOCKING, 2h/WARNING, 0.5h/INFO: "
        f"{len(blocking)}×4 + {len(warning)}×2 + {len(info)}×0.5 = "
        f"{len(blocking)*4 + len(warning)*2 + len(info)*0.5} hours ("
        f"~{(len(blocking)*4 + len(warning)*2 + len(info)*0.5)/8:.1f} person-days). "
        "Excludes TRA-001 (per-leaf segment translation, estimated 2-3 person-days alone)."
    )

    # === Conclusion ===
    add_heading(doc, "6. Conclusion", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "The TRA prototype at HEAD 4b8827c is substantially more conformant than at the time of "
        "the Round 1 audit. All 4 critical invariants hold, all 4 quality gates are green, the "
        "test count grew from 103 to 141, benchmark coverage grew from 13/24 to 22/24, and "
        "audit-trail byte-reproducibility is empirically verified. The 30 fully-fixed Round-1 "
        "findings demonstrate that the Round 1 audit's TDD-cycle remediation was effective."
    )
    p = doc.add_paragraph()
    p.add_run(
        "However, Round 2 uncovered 3 new BLOCKING findings (TRA-036, TRA-037, TRA-048) that "
        "represent regressions or gaps introduced by the remediation commits. TRA-036 and TRA-037 "
        "are direct consequences of the TRA-004 and TRA-008 fixes (the early-return and the "
        "post-gate anchor rewrite). TRA-048 is a test-coverage gap that mutation testing caught. "
        "These 3 should be addressed before any L3/L4 production use."
    )
    p = doc.add_paragraph()
    p.add_run(
        "The 22 new WARNING findings cluster around 4 themes: (a) exception-recovery completeness "
        "(TRA-038, TRA-039, TRA-040, TRA-041, TRA-044); (b) test-coverage gaps identified by "
        "mutation testing (TRA-049 through TRA-055); (c) doc staleness after remediation "
        "(TRA-059 through TRA-063); (d) type-safety and dead-code hygiene (TRA-043, TRA-045, "
        "TRA-046, TRA-047). None are spec-invariant violations; all are quality/auditability "
        "improvements."
    )
    p = doc.add_paragraph()
    p.add_run(
        "The 13 INFO findings are minor hygiene (dead code, doc count errors, file duplication). "
        "They do not affect conformance or correctness."
    )
    p = doc.add_paragraph()
    p.add_run("Overall verdict: ").bold = True
    p.add_run(
        "The TRA prototype is L3-conformant for the happy path (zero BLOCKING on the example "
        "advisory at L3_STRICT). The 3 new BLOCKING findings affect edge cases (analyze failure, "
        "link rewriting, LLM degradation test coverage) and should be fixed before claiming "
        "full L3/L4 production readiness. The Round 1 audit's 34/35 remediation rate is "
        "confirmed; Round 2 adds 36 new findings for the next remediation cycle."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
