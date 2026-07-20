"""Generate formal DOCX audit report for Round 7.

Output: /home/z/my-project/Translation-Runtime-Architecture/docs/audit/round7/TRA_Prototype_Audit_Report_r7.docx
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REGISTER = Path(
    "/home/z/my-project/Translation-Runtime-Architecture/docs/audit/round7/master_findings_register_r7.json"
)
HEATMAP = Path(
    "/home/z/my-project/Translation-Runtime-Architecture/docs/audit/round7/TRA_audit_severity_heatmap_r7.png"
)
OUT = Path(
    "/home/z/my-project/Translation-Runtime-Architecture/docs/audit/round7/TRA_Prototype_Audit_Report_r7.docx"
)

SEVERITY_COLORS = {
    "BLOCKING": RGBColor(0xCC, 0x00, 0x00),
    "WARNING": RGBColor(0xCC, 0x88, 0x00),
    "INFO": RGBColor(0x00, 0x66, 0x00),
}


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x20, 0x40, 0x60)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_finding(doc: Document, finding: dict) -> None:
    sev = finding["severity"]
    color = SEVERITY_COLORS.get(sev, RGBColor(0, 0, 0))

    h = doc.add_heading(level=2)
    run_id = h.add_run(f"{finding['id']}: ")
    run_id.bold = True
    run_id.font.size = Pt(13)
    run_title = h.add_run(finding["title"])
    run_title.font.size = Pt(13)
    run_sev = h.add_run(f"  [{sev}]")
    run_sev.bold = True
    run_sev.font.color.rgb = color
    run_sev.font.size = Pt(11)

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Light Grid Accent 1"
    cells = [
        ("Track(s)", finding["track"]),
        ("Category", finding["category"]),
        ("Round 6 Status", finding.get("round6_status", "new") or "new"),
        ("Finding Type", finding.get("finding_type", "issue")),
        ("Root ID", finding.get("root_id", finding["id"])),
    ]
    for i, (k, v) in enumerate(cells):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = str(v)[:200]
        for para in tbl.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.bold = True

    p = doc.add_paragraph()
    p.add_run("Evidence: ").bold = True
    p.add_run(finding.get("evidence", "(none)"))

    p = doc.add_paragraph()
    p.add_run("Detail: ").bold = True
    p.add_run(finding.get("detail", "(none)"))

    p = doc.add_paragraph()
    p.add_run("Suggested fix: ").bold = True
    p.add_run(finding.get("suggested_fix", "(none)"))

    doc.add_paragraph()


def main() -> None:
    findings = json.loads(REGISTER.read_text(encoding="utf-8"))
    issues = [f for f in findings if f.get("finding_type") == "issue"]
    positives = [f for f in findings if f.get("finding_type") == "positive_verification"]
    by_sev: dict[str, int] = {"BLOCKING": 0, "WARNING": 0, "INFO": 0}
    for f in issues:
        by_sev[f["severity"]] = by_sev.get(f["severity"], 0) + 1

    by_r6_status: dict[str, int] = {}
    for f in findings:
        s = f.get("round6_status", "new") or "new"
        sl = s.lower()
        if "fixed" in sl and "verified" in sl:
            key = "fixed-and-verified"
        elif "fixed" in sl:
            key = "fixed"
        elif "partial" in sl:
            key = "partial"
        elif "persistent" in sl:
            key = "persistent"
        elif "regression" in sl:
            key = "new-regression"
        elif "new" in sl:
            key = "new"
        else:
            key = "other"
        by_r6_status[key] = by_r6_status.get(key, 0) + 1

    doc = Document()

    # Cover
    title = doc.add_heading("TRA Prototype Engine", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Round 7 Independent Re-Audit Report", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HEAD audited: ").bold = True
    p.add_run("6d3144a3fdaa8d90a8f5b5f3996af39e667ee496\n")
    p.add_run("Audit date: ").bold = True
    p.add_run("2026-07-21\n")
    p.add_run("Methodology: ").bold = True
    p.add_run("7-track parallel re-audit (R7/A7/B7/C7/D7/E7/F7)\n")
    p.add_run("Carry-over input: ").bold = True
    p.add_run("Round 6 master register (76 findings: 58 issues + 18 positive verifications)\n")
    p.add_run("Auditor: ").bold = True
    p.add_run("Super Z (independent)")

    doc.add_paragraph()

    # Executive Summary
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        f"Round 7 is the sixth independent re-audit of the TRA prototype engine, "
        f"conducted at HEAD 6d3144a — one commit past the Round 6 baseline (c4ecd41). "
        f"That commit (\"fix(tra): Round 6 Batch 1 — fix 4 BLOCKING + 6 WARNING "
        f"findings via TDD\") remediated all 4 R6 BLOCKING doc-staleness findings "
        f"and 6 of the 10 R6 WARNING findings. Round 7 verifies the R6 Batch 1 "
        f"fixes hold, hunts for NEW issues introduced by the remediation, and "
        f"re-confirms the 18 R6 positive verifications.",
    )
    add_para(
        doc,
        f"The audit produced {len(findings)} deduplicated findings: "
        f"{len(issues)} issues and {len(positives)} positive verifications. "
        f"Of the issues, {by_sev['BLOCKING']} are BLOCKING, "
        f"{by_sev['WARNING']} are WARNING, and {by_sev['INFO']} are INFO. "
        f"Track R7 re-verified all 76 Round 6 entries (58 issues + 18 positives) "
        f"at HEAD 6d3144a.",
    )
    add_para(
        doc,
        f"All 4 critical invariants hold at HEAD 6d3144a with code-level evidence. "
        f"The 3 OWASP security fixes (TRA-076/077/078) remain effective. TRA-013 "
        f"byte-reproducibility is re-verified within HEAD. Round 6 Batch 1 fix "
        f"(audit_trace.jsonl truncate mode in TRAKernel) closes the E6-001 "
        f"reproducibility gap that previously manifested on reused CLI default paths.",
    )

    # Methodology
    add_heading(doc, "Methodology", level=1)
    add_para(
        doc,
        "Round 7 followed the same 7-track parallel structure refined through "
        "Rounds 3-6. Each track operated independently against HEAD 6d3144a. "
        "Each track produced a Markdown findings file with a uniform structure: "
        "summary, findings (each with severity, category, evidence, detail, "
        "suggested fix, Round 6 status), Round 6 carry-over matrix, and "
        "verification commands. Findings were synthesized into a master register "
        "via synthesize_findings_r7.py, with deduplication by root finding ID "
        "(TRA-NNN).",
    )
    add_para(
        doc,
        "Severity lexicon mirrors the TRA spec: BLOCKING (must fix before L3 "
        "certification), WARNING (should fix, does not block certification), INFO "
        "(cosmetic or minor). Each finding cites at least 2 file:line evidence "
        "points at HEAD 6d3144a.",
    )
    add_para(
        doc,
        "Quality gates were re-run at HEAD: ruff format (clean), ruff check "
        "(clean), mypy --strict (0 issues, 20 source files), pytest (309 passed). "
        "TRA-013 byte-reproducibility was independently re-verified by Track E7: "
        "two cold-cache L4 translations of to_translate.md produce byte-identical "
        "audit_trace.jsonl, evidence_trace.jsonl, and output markdown.",
    )

    # Track summaries
    add_heading(doc, "Track Summaries", level=1)

    def track_counts(prefix: str) -> tuple[int, int, int]:
        b = w = i = 0
        for f in issues:
            if prefix in f["track"].split(","):
                if f["severity"] == "BLOCKING":
                    b += 1
                elif f["severity"] == "WARNING":
                    w += 1
                else:
                    i += 1
        return b, w, i

    a7b, a7w, a7i = track_counts("A7")
    b7b, b7w, b7i = track_counts("B7")
    c7b, c7w, c7i = track_counts("C7")
    d7b, d7w, d7i = track_counts("D7")
    e7b, e7w, e7i = track_counts("E7")
    f7b, f7w, f7i = track_counts("F7")

    track_summaries = [
        (
            "Track R7 — Regression Baseline",
            f"Re-verified all 76 Round 6 entries (58 issues + 18 positive "
            f"verifications) at HEAD 6d3144a. Round 6 Batch 1 remediation is "
            f"verified landed: 4 BLOCKING doc-staleness fixes (TRA-C6-001/002/006/"
            f"D6-001) + 6 WARNING fixes (mutmut config, cache-isolation, vacuous "
            f"HITL tests, audit_trace truncate mode). See R6 Status sheet in the "
            f"XLSX register for the full carry-over matrix.",
        ),
        (
            "Track A7 — Spec Conformance",
            f"{a7b + a7w + a7i} findings ({a7b} BLOCKING / {a7w} WARNING / {a7i} INFO). "
            f"4 critical invariants all hold at HEAD with code-level evidence: "
            f"canonical terminology exact, entities immutable, VERIFY_OUTPUT never "
            f"self-scores, REPAIR_SEGMENT surgical. All 5 PolicyResolver severity "
            f"pairs arbitrated. Per-leaf translation (TRA-001 Phase 8) works.",
        ),
        (
            "Track B7 — Code Quality & Security",
            f"{b7b + b7w + b7i} findings ({b7b} BLOCKING / {b7w} WARNING / {b7i} INFO). "
            f"3 OWASP fixes verified holding (TRA-076 A03, TRA-077 A08, TRA-078 A09). "
            f"TRA-013 byte-identical L4 reproducibility confirmed within HEAD. "
            f"All 4 quality gates green (309 tests).",
        ),
        (
            "Track C7 — Doc-vs-Code Consistency",
            f"{c7b + c7w + c7i} findings ({c7b} BLOCKING / {c7w} WARNING / {c7i} INFO). "
            f"R6 Batch 1 refreshed the worst doc drift (CLAUDE.md, README.md, "
            f"SKILL.md, implementation_plan.md, to_translate.md). Verify whether "
            f"all stale claims are now accurate.",
        ),
        (
            "Track D7 — Test Suite",
            f"{d7b + d7w + d7i} findings ({d7b} BLOCKING / {d7w} WARNING / {d7i} INFO). "
            f"Test count: 309 across 16 test files. Benchmark cases: 36 "
            f"(S/F/T/D/E/R coverage). R6 Batch 1 fixed cache-pollution, mutmut "
            f"config, vacuous HITL tests, wrong field name in model_copy.",
        ),
        (
            "Track E7 — Forensic L4 End-to-End",
            f"{e7b + e7w + e7i} findings ({e7b} BLOCKING / {e7w} WARNING / {e7i} INFO). "
            f"TRA-013 byte-reproducibility HOLDS within HEAD — 2 cold-cache L4 runs "
            f"produce byte-identical audit_trace.jsonl + evidence_trace.jsonl. "
            f"9/9 expected L4 artifacts present + valid. R6 Batch 1 truncate mode "
            f"closes the append-mode reproducibility gap on reused CLI default paths.",
        ),
        (
            "Track F7 — Stub-Module Conformance",
            f"{f7b + f7w + f7i} findings ({f7b} BLOCKING / {f7w} WARNING / {f7i} INFO). "
            f"TRA-096 (as_interface), TRA-099 (CLI --registry), TRA-F5-010 "
            f"(normalize_language_pair), TRA-F5-011 (register direction check) all "
            f"verified holding. ModuleInterface contract matches "
            f"LanguageModuleProtocol.",
        ),
    ]
    for title_str, body in track_summaries:
        add_heading(doc, title_str, level=2)
        add_para(doc, body)

    # Heatmap
    if HEATMAP.exists():
        add_heading(doc, "Severity Heatmap", level=1)
        doc.add_picture(str(HEATMAP), width=Inches(6.5))
        doc.add_paragraph()

    # Findings
    add_heading(doc, "Findings Register (Issues Only)", level=1)
    add_para(
        doc,
        f"{len(issues)} issues, sorted by severity (BLOCKING first). "
        f"{len(positives)} positive verifications (invariants holding, fixes "
        f"confirmed) are omitted from this section for brevity; they are listed in "
        f"the master register JSON.",
        italic=True,
    )

    sev_order = ["BLOCKING", "WARNING", "INFO"]
    for sev in sev_order:
        sev_findings = [f for f in issues if f["severity"] == sev]
        if not sev_findings:
            continue
        add_heading(doc, f"{sev} Findings ({len(sev_findings)})", level=2)
        for f in sev_findings:
            add_finding(doc, f)

    # Conclusion
    add_heading(doc, "Conclusion", level=1)
    add_para(
        doc,
        f"HEAD 6d3144a is conformant to TRA spec §1–§9 at the level of the 4 "
        f"critical invariants and the 6 ISA instruction contracts. The single "
        f"Round 6 Batch 1 commit successfully remediated all 4 R6 BLOCKING "
        f"doc-staleness findings and 6 of the 10 R6 WARNING findings. No "
        f"regressions were introduced. The 3 OWASP security fixes (TRA-076/077/078) "
        f"remain effective. TRA-013 byte-reproducibility is maintained within HEAD.",
    )
    add_para(
        doc,
        f"Test count held at 309 across 16 test files (R6 Batch 1 modified tests "
        f"in place rather than adding new ones). Benchmark suite remains at 36 "
        f"cases. The remaining WARNING findings (TRA-A6-001 cache-hit suppresses "
        f"EXCEPTION_HANDLER, TRA-A6-002 segment_index plumbing, TRA-D6-003 "
        f"hardcoded paths, TRA-D6-004 vacuous entity-ambiguity test) plus 44 INFO "
        f"findings are addressed in remediation_plan_r7.md.",
    )
    add_para(
        doc,
        "Recommendation: address the remaining WARNING cluster (TRA-A6-001 + "
        "TRA-A6-002 + TRA-D6-003 + TRA-D6-004) as a single forensic-completeness "
        "sprint, since all four touch the per-leaf translation / cache-hit / "
        "audit-record emission surface introduced by R5 Batch 2 + H. See "
        "remediation_plan_r7.md for the TDD plan and the Remediation Backlog "
        "sheet in the XLSX register for effort estimates.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote: {OUT}")
    print(f"  Findings: {len(findings)} total ({len(issues)} issues + {len(positives)} positives)")
    print(f"  Issues by severity: {by_sev}")


if __name__ == "__main__":
    main()
