"""Generate formal DOCX audit report for Round 4.

Output: /home/z/my-project/download/TRA_Round4/TRA_Prototype_Audit_Report_r4.docx
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REGISTER = Path(
    "/home/z/my-project/Translation-Runtime-Architecture/docs/audit/round4/master_findings_register_r4.json"
)
HEATMAP = Path(
    "/home/z/my-project/download/TRA_Round4/TRA_audit_severity_heatmap_r4.png"
)
OUT = Path(
    "/home/z/my-project/download/TRA_Round4/TRA_Prototype_Audit_Report_r4.docx"
)

SEVERITY_COLORS = {
    "BLOCKING": RGBColor(0xCC, 0x00, 0x00),
    "WARNING": RGBColor(0xCC, 0x88, 0x00),
    "INFO": RGBColor(0x00, 0x66, 0x00),
}


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x20, 0x40, 0x60)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_finding(doc: Document, finding: dict) -> None:
    sev = finding["severity"]
    color = SEVERITY_COLORS.get(sev, RGBColor(0, 0, 0))

    # Heading: ID — Title (Severity)
    h = doc.add_heading(level=2)
    run_id = h.add_run(f"{finding['id']}: ")
    run_id.bold = True
    run_id.font.size = Pt(13)
    run_title = h.add_run(finding["title"])
    run_title.font.size = Pt(13)
    run_sev = h.add_run(f"  [{sev}]")
    run_sev.bold = True
    run_sev.font.color.rgb = color
    run_sev.font.size = Pt(11)

    # Metadata table
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Light Grid Accent 1"
    cells = [
        ("Track(s)", finding["track"]),
        ("Category", finding["category"]),
        ("Round 3 Status", finding.get("round3_status", "new") or "new"),
        ("Finding Type", finding.get("finding_type", "issue")),
        ("Root ID", finding.get("root_id", finding["id"])),
    ]
    for i, (k, v) in enumerate(cells):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = str(v)[:200]
        for para in tbl.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.bold = True

    # Evidence
    p = doc.add_paragraph()
    p.add_run("Evidence: ").bold = True
    p.add_run(finding.get("evidence", "(none)"))

    # Detail
    p = doc.add_paragraph()
    p.add_run("Detail: ").bold = True
    p.add_run(finding.get("detail", "(none)"))

    # Suggested fix
    p = doc.add_paragraph()
    p.add_run("Suggested fix: ").bold = True
    p.add_run(finding.get("suggested_fix", "(none)"))

    doc.add_paragraph()  # spacer


def main() -> None:
    findings = json.loads(REGISTER.read_text(encoding="utf-8"))
    issues = [f for f in findings if f.get("finding_type") == "issue"]
    positives = [f for f in findings if f.get("finding_type") == "positive_verification"]
    by_sev = {"BLOCKING": 0, "WARNING": 0, "INFO": 0}
    for f in issues:
        by_sev[f["severity"]] = by_sev.get(f["severity"], 0) + 1

    doc = Document()

    # ----- Cover -----
    title = doc.add_heading("TRA Prototype Engine", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Round 4 Independent Re-Audit Report", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HEAD audited: ").bold = True
    p.add_run("805a8f8c9843cd429b30623a1a84b336b7920e4c\n")
    p.add_run("Audit date: ").bold = True
    p.add_run("2026-07-17\n")
    p.add_run("Methodology: ").bold = True
    p.add_run("7-track parallel re-audit (R4/A4/B4/C4/D4/E4/F4)\n")
    p.add_run("Carry-over input: ").bold = True
    p.add_run("Round 3 master register (36 findings)\n")
    p.add_run("Auditor: ").bold = True
    p.add_run("Super Z (independent)")

    doc.add_paragraph()

    # ----- Executive Summary -----
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        f"Round 4 is the fourth independent re-audit of the TRA prototype engine. "
        f"It was conducted at HEAD 805a8f8, six commits past the Round 3 baseline "
        f"(b783745). The audit used a 7-track parallel structure extending Round 3's "
        f"methodology: Track R4 (regression baseline), Track A4 (spec conformance), "
        f"Track B4 (code quality & security), Track C4 (doc-vs-code consistency), "
        f"Track D4 (test suite), Track E4 (forensic L4 end-to-end), and Track F4 "
        f"(stub-module conformance).",
    )
    add_para(
        doc,
        f"The audit produced {len(findings)} deduplicated findings: "
        f"{len(issues)} issues and {len(positives)} positive verifications. "
        f"Of the issues, {by_sev['BLOCKING']} is BLOCKING, "
        f"{by_sev['WARNING']} are WARNING, and {by_sev['INFO']} are INFO. "
        f"The 19 positive verifications confirm that the 4 critical invariants hold, "
        f"the 3 OWASP security fixes (TRA-076/077/078) are still effective, and "
        f"TRA-013 byte-reproducibility is maintained (audit_trace.jsonl sha256 "
        f"263b901e..., matches R3 exactly).",
    )
    add_para(
        doc,
        f"Track R4 re-verified all 36 Round 3 findings: 20 FIXED, 12 PERSISTENT, "
        f"4 PARTIAL, 0 REGRESSED. The Round 3 BLOCKING findings (TRA-093 + TRA-096) "
        f"are both verified fixed at HEAD. The single Round 4 BLOCKING finding "
        f"(TRA-C4-013) is a documentation defect: tra-prototype/README.md's "
        f"'Commands' section uses bare tra_cli.py invocations that fail because the "
        f"file is not executable, has no shebang, and is not on PATH. This is a "
        f"user-facing onboarding break, not a code defect.",
    )
    add_para(
        doc,
        "The most material drift this round is in documentation: CLAUDE.md, "
        "SKILL.md, and tra-prototype/README.md all mark TRA-017 (unused deps) as "
        "'persistent' when it is in fact FIXED (commit a3cd2c1). SKILL.md §7 claims "
        "'174 tests across 16 test files' but the actual count is 199 across 18. "
        "These drifts are tracked in Track C4 (17 findings, 1 BLOCKING).",
    )

    # ----- Methodology -----
    add_heading(doc, "Methodology", level=1)
    add_para(
        doc,
        "Round 4 followed the same 7-track parallel structure introduced in Round 3, "
        "with each track operating independently against the same HEAD (805a8f8). "
        "Each track produced a Markdown findings file with a uniform structure: "
        "summary, findings (each with severity, category, evidence, detail, "
        "suggested fix, Round 3 status), Round 3 carry-over matrix, and "
        "verification commands. Findings were synthesized into a master register "
        "via a Python script (synthesize_findings_r4.py), with deduplication by "
        "root finding ID (TRA-NNN).",
    )
    add_para(
        doc,
        "Severity lexicon mirrors the TRA spec: BLOCKING (must fix before L3 "
        "certification), WARNING (should fix, does not block certification), INFO "
        "(cosmetic or minor). No escalation or de-escalation from Round 3 without "
        "explicit justification. Each finding cites at least 2 file:line evidence "
        "points at HEAD 805a8f8.",
    )
    add_para(
        doc,
        "Quality gates were re-run at HEAD: ruff format (40 files), ruff check "
        "(clean), mypy --strict (0 issues), pytest (199 passed). TRA-013 "
        "byte-reproducibility was independently re-verified by Track E4: two "
        "cold-cache L4 translations of the same source produced byte-identical "
        "audit_trace.jsonl, evidence_trace.jsonl, and output markdown (sha256 "
        "hashes match R3 exactly).",
    )

    # ----- Track summaries -----
    add_heading(doc, "Track Summaries", level=1)

    track_summaries = [
        (
            "Track R4 — Regression Baseline",
            "Re-verified all 36 Round 3 findings via regression test or static "
            "check. Result: 20 FIXED, 12 PERSISTENT, 4 PARTIAL, 0 REGRESSED. "
            "Notable: TRA-017 deps trimmed (fixed); TRA-038 exception routable "
            "but auto-detection deferred (partial); TRA-099 CLI --registry still "
            "not wired (persistent); TRA-085 status.md banner itself stale (says "
            "174+, actual is 199).",
        ),
        (
            "Track A4 — Spec Conformance",
            "11 findings (0 BLOCKING / 6 WARNING / 5 INFO). 4 critical invariants "
            "all hold at HEAD with code-level evidence: canonical terminology "
            "exact, entities immutable, VERIFY_OUTPUT never self-scores, "
            "REPAIR_SEGMENT surgical. 1 NEW finding: TRA-A4-011 (repaired = "
            "repaired no-op self-assignment at isa.py:654, parallel to TRA-073, "
            "missed by R3).",
        ),
        (
            "Track B4 — Code Quality & Security",
            "17 findings (0 BLOCKING / 3 WARNING / 14 INFO). 3 OWASP fixes "
            "verified holding (TRA-076 A03, TRA-077 A08, TRA-078 A09) — "
            "mutation-tested. TRA-013 byte-identical L4 reproducibility "
            "confirmed: audit_trace.jsonl sha256 263b901e..., evidence_trace.jsonl "
            "sha256 f9831523.... All 4 quality gates green.",
        ),
        (
            "Track C4 — Doc-vs-Code Consistency",
            "17 findings (1 BLOCKING / 9 WARNING / 7 INFO). The highest-value "
            "track this round. 1 BLOCKING: tra-prototype/README.md 'Commands' "
            "section uses bare tra_cli.py invocations — file is mode 664 (not "
            "executable), no shebang, no [project.scripts] entry point. All 4 CLI "
            "examples fail as written. 12 new findings include: CLAUDE.md + "
            "tra-prototype/README.md both mark TRA-017 as persistent but it's "
            "FIXED; SKILL.md §7 says '174 tests' (actual 199, 14.4% drift); "
            "SKILL.md §8 lists TRA-016/017/026 as unfixed but all 3 are FIXED.",
        ),
        (
            "Track D4 — Test Suite",
            "15 findings (0 BLOCKING / 5 WARNING / 10 INFO). Test count: 199 "
            "(R3: 174, +25 — all in test_outstanding_findings.py, 12 new "
            "TestTRA0XX classes). Benchmark cases: 22 (S:5/F:5/T:5/D:4/E:2/R:1) "
            "— still 22% of spec's 100+ target, no growth across 3 rounds. S-03 "
            "and E-03 still missing. New finding TRA-D4-014: redundant 186-LOC "
            "manual e2e script (run_e2e_translation.py added in commit 805a8f8) "
            "duplicates tests/test_e2e_to_translate.py.",
        ),
        (
            "Track E4 — Forensic L4 End-to-End",
            "15 findings (0 BLOCKING / 2 WARNING / 13 INFO). TRA-013 "
            "byte-reproducibility HOLDS — all 6 sha256 hashes match across 2 "
            "cold-cache runs. 9/9 expected L4 artifacts present + valid. R3 "
            "BLOCKING TRA-E3-003 (false-positive BROKEN_LINK on CJK headings) "
            "fully resolved by TRA-093 fix. 1 new INFO finding: style_profile.yaml "
            "not documented in SKILL.md §4.",
        ),
        (
            "Track F4 — Stub-Module Conformance",
            "7 findings (0 BLOCKING / 2 WARNING / 5 INFO). TRA-096 (as_interface) "
            "FIXED — verified end-to-end via a stub fr-en module test. TRA-099 "
            "(CLI --registry) PERSISTENT — CLI still constructs TRAKernel without "
            "registry, silently overrides --lang fr-en with ZHENModule fallback. "
            "2 new findings: TRA-F4-006 (minimal ModuleInterface default-mismatch "
            "crash), TRA-F4-007 (_select_module silent dispatch on same-source-lang "
            "collisions).",
        ),
    ]
    for title, body in track_summaries:
        add_heading(doc, title, level=2)
        add_para(doc, body)

    # ----- Heatmap -----
    if HEATMAP.exists():
        add_heading(doc, "Severity Heatmap", level=1)
        doc.add_picture(str(HEATMAP), width=Inches(6.5))
        doc.add_paragraph()

    # ----- Findings -----
    add_heading(doc, "Findings Register (Issues Only)", level=1)
    add_para(
        doc,
        f"{len(issues)} issues, sorted by severity (BLOCKING first). "
        f"19 positive verifications (invariants holding, fixes confirmed) are "
        f"omitted from this section for brevity; they are listed in the master "
        f"register JSON.",
        italic=True,
    )

    sev_order = ["BLOCKING", "WARNING", "INFO"]
    for sev in sev_order:
        sev_findings = [f for f in issues if f["severity"] == sev]
        if not sev_findings:
            continue
        add_heading(doc, f"{sev} Findings ({len(sev_findings)})", level=2)
        for f in sev_findings:
            add_finding(doc, f)

    # ----- Conclusion -----
    add_heading(doc, "Conclusion", level=1)
    add_para(
        doc,
        f"HEAD 805a8f8 is conformant to TRA spec §1–§9 at the level of the 4 "
        f"critical invariants and the 6 ISA instruction contracts. The 6 commits "
        f"since Round 3 successfully remediated the 2 Round 3 BLOCKING findings "
        f"(TRA-093, TRA-096) plus 18 WARNING/INFO findings. No regressions were "
        f"introduced. The 3 OWASP security fixes (TRA-076/077/078) remain "
        f"effective, mutation-tested. TRA-013 byte-reproducibility is maintained.",
    )
    add_para(
        doc,
        f"The single Round 4 BLOCKING finding (TRA-C4-013) is a documentation "
        f"defect in tra-prototype/README.md — the 'Commands' section's CLI "
        f"invocation form is broken. This is an onboarding break, not a code "
        f"defect; the root README.md and SKILL.md correctly use 'python -m "
        f"tra_cli'. The fix is a 4-line README edit.",
    )
    add_para(
        doc,
        f"The 11 WARNING findings are dominated by documentation drift (Track C4) "
        f"and persistent Phase 0 prototype gaps (TRA-001 whole-doc translation, "
        f"TRA-038 unreachable exceptions, TRA-042 structural verification, "
        f"TRA-072 universal Policy Engine arbitration, TRA-099 CLI --registry). "
        f"None represent a regression from Round 3; all are documented in "
        f"SKILL.md's 'Known limitations' section (though that section is itself "
        f"drifted — see Track C4).",
    )
    add_para(
        doc,
        "Recommendation: prioritize the BLOCKING doc fix (TRA-C4-013) and the "
        "Track C4 doc-staleness findings as a single documentation-refresh batch. "
        "The TRA-001 / TRA-038 / TRA-072 / TRA-042 cluster is a larger spec-"
        "conformance effort that should be planned as a separate Phase 8 "
        "remediation sprint. See the Remediation Backlog sheet in the XLSX "
        "register for effort estimates.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote: {OUT}")
    print(f"  Findings: {len(findings)} total ({len(issues)} issues + {len(positives)} positives)")
    print(f"  Issues by severity: {by_sev}")


if __name__ == "__main__":
    main()
